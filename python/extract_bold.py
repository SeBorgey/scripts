from docx import Document

def extract_bold_text(input_file, output_file):
    doc = Document(input_file)
    new_doc = Document()

    for para in doc.paragraphs:
        bold_text = ""
        for run in para.runs:
            if run.bold:
                bold_text += run.text

        if bold_text.strip():
            new_doc.add_paragraph(bold_text)

    new_doc.save(output_file)

input_file = "3.docx"
output_file = "output3.docx"

extract_bold_text(input_file, output_file)
